#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import subprocess
from datetime import datetime
from PySide2.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QTreeView, QFileSystemModel,
    QSplitter, QGroupBox, QGridLayout, QComboBox, 
    QMessageBox, QStatusBar, QFileDialog, QTabWidget, QTextEdit
)
from PySide2.QtCore import Qt, QDir, QSettings
from PySide2.QtGui import QScreen

# 尝试导入 python-docx 用于生成符合国标的 Word 文档
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class GenDirApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("GenDir Qt-Arm-v0.08 - 高级国标公文与目录生成器")
        
        # 1. 根据屏幕分辨率自适应调整界面大小为屏幕的 85% 并居中显示
        screen = QApplication.primaryScreen()
        screen_rect = screen.availableGeometry()
        
        app_width = int(screen_rect.width() * 0.85)
        app_height = int(screen_rect.height() * 0.85)
        self.resize(app_width, app_height)
        
        # 居中显示
        center_x = screen_rect.x() + (screen_rect.width() - app_width) // 2
        center_y = screen_rect.y() + (screen_rect.height() - app_height) // 2
        self.move(center_x, center_y)
        
        # 初始化设置
        self.settings = QSettings("GenDir", "QtArm")
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)
        
        # 创建左右分栏 (1:1 相同宽度)
        splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(splitter)
        
        # ========== 左侧面板：路径导航 + 时间与文件夹创建控件 ==========
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        splitter.addWidget(left_widget)
        
        # 1. 全盘路径导航
        nav_group = QGroupBox("全盘路径导航")
        nav_layout = QVBoxLayout(nav_group)
        
        path_layout = QHBoxLayout()
        
        # 增加“上一层”按钮
        btn_parent = QPushButton("⬆️ 上一层")
        btn_parent.setToolTip("返回上一层目录")
        btn_parent.clicked.connect(self.go_parent_path)
        path_layout.addWidget(btn_parent)
        
        self.path_edit = QLineEdit()
        self.path_edit.setPlaceholderText("输入路径或点击右侧按钮浏览...")
        self.path_edit.textChanged.connect(self.on_path_changed)
        path_layout.addWidget(self.path_edit)
        
        btn_browse = QPushButton("📁 浏览...")
        btn_browse.clicked.connect(self.browse_path)
        path_layout.addWidget(btn_browse)
        
        btn_set_path = QPushButton("设为当前")
        btn_set_path.clicked.connect(self.set_current_path)
        path_layout.addWidget(btn_set_path)
        nav_layout.addLayout(path_layout)
        
        # 文件系统树
        self.tree_view = QTreeView()
        self.tree_view.setMinimumHeight(220)
        self.file_model = QFileSystemModel()
        self.file_model.setRootPath(QDir.rootPath())
        self.file_model.setFilter(QDir.AllEntries | QDir.NoDotAndDotDot)
        self.tree_view.setModel(self.file_model)
        self.tree_view.setRootIndex(self.file_model.index(QDir.rootPath()))
        self.tree_view.doubleClicked.connect(self.on_tree_double_click)
        nav_layout.addWidget(self.tree_view)
        left_layout.addWidget(nav_group)
        
        # 2. 时间格式设置
        time_group = QGroupBox("时间格式设置")
        time_layout = QHBoxLayout(time_group)
        time_layout.addWidget(QLabel("时间戳格式："))
        self.time_format_combo = QComboBox()
        self.time_format_combo.addItems([
            "标准点分格式 (YYYY.MM.DD)",
            "精细时间格式 (YYYY.MM.DD HHMMSS)"
        ])
        self.time_format_combo.currentIndexChanged.connect(self.update_preview)
        time_layout.addWidget(self.time_format_combo)
        left_layout.addWidget(time_group)
        
        # 3. 创建文件夹与子目录
        folder_group = QGroupBox("创建文件夹与子目录")
        folder_layout = QVBoxLayout(folder_group)
        
        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel("主文件夹名："))
        self.folder_name_edit = QLineEdit()
        self.folder_name_edit.setPlaceholderText("输入主文件夹基础名称")
        self.folder_name_edit.textChanged.connect(self.update_preview)
        name_layout.addWidget(self.folder_name_edit)
        folder_layout.addLayout(name_layout)
        
        sub_layout = QHBoxLayout()
        sub_layout.addWidget(QLabel("子文件夹列表："))
        self.sub_folder_edit = QLineEdit()
        self.sub_folder_edit.setPlaceholderText("多个子文件夹用逗号分隔，例: 01材料,02总结,03附件")
        sub_layout.addWidget(self.sub_folder_edit)
        folder_layout.addLayout(sub_layout)
        
        create_btn_layout = QHBoxLayout()
        btn_create = QPushButton("📁 仅创建主文件夹")
        btn_create.clicked.connect(self.create_folder)
        create_btn_layout.addWidget(btn_create)
        
        btn_create_sub = QPushButton("📁 创建主夹及子文件夹")
        btn_create_sub.clicked.connect(self.create_sub_folders)
        create_btn_layout.addWidget(btn_create_sub)
        folder_layout.addLayout(create_btn_layout)
        
        # 文件夹命名预览
        preview_layout = QHBoxLayout()
        preview_layout.addWidget(QLabel("文件夹命名预览:"))
        self.preview_label = QLabel("新文件夹 2026.07.25")
        self.preview_label.setStyleSheet("""
            QLabel {
                background-color: #f8f9fa;
                padding: 4px 8px;
                border: 1px solid #ced4da;
                border-radius: 4px;
                font-family: monospace;
                color: #333;
            }
        """)
        preview_layout.addWidget(self.preview_label)
        folder_layout.addLayout(preview_layout)
        
        left_layout.addWidget(folder_group)
        
        # ========== 右侧面板：Tab页面 (公文生成 / 其他文件生成 / 档案封面生成) ==========
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        splitter.addWidget(right_widget)
        
        # 当前目标路径显示
        target_path_layout = QHBoxLayout()
        target_path_layout.addWidget(QLabel("目标路径:"))
        self.file_path_edit = QLineEdit()
        self.file_path_edit.setReadOnly(True)
        target_path_layout.addWidget(self.file_path_edit)
        right_layout.addLayout(target_path_layout)
        
        # 创建 TabWidget
        self.tab_widget = QTabWidget()
        right_layout.addWidget(self.tab_widget)
        
        # --- Tab 1: 公文生成 (国标 docx) ---
        tab_doc = QWidget()
        tab_doc_layout = QVBoxLayout(tab_doc)
        
        doc_form_layout = QGridLayout()
        doc_form_layout.addWidget(QLabel("公文标题/文号："), 0, 0)
        self.doc_title_edit = QLineEdit()
        self.doc_title_edit.setPlaceholderText("请输入公文标题或字号")
        doc_form_layout.addWidget(self.doc_title_edit, 0, 1)
        
        doc_form_layout.addWidget(QLabel("密级与紧急程度："), 1, 0)
        self.doc_level_combo = QComboBox()
        self.doc_level_combo.addItems(["普通", "内部", "秘密", "机密", "加急", "特急"])
        doc_form_layout.addWidget(self.doc_level_combo, 1, 1)
        tab_doc_layout.addLayout(doc_form_layout)
        
        tab_doc_layout.addWidget(QLabel("公文预置内容细节 (支持 '一、'、'（一）' 自动识别排版)："))
        self.doc_content_edit = QTextEdit()
        self.doc_content_edit.setPlaceholderText("在此输入公文正文细节或模板框架...")
        self.doc_content_edit.setPlainText("一、 工作进展情况...\n二、 下一步工作安排...")
        tab_doc_layout.addWidget(self.doc_content_edit)
        
        btn_gen_doc = QPushButton("📄 生成国标公文 Word 文档 (.docx)")
        btn_gen_doc.clicked.connect(self.generate_document)
        tab_doc_layout.addWidget(btn_gen_doc)
        
        self.tab_widget.addTab(tab_doc, "公文生成")
        
        # --- Tab 2: 其他文件生成 ---
        tab_other = QWidget()
        tab_other_layout = QVBoxLayout(tab_other)
        
        other_form_layout = QGridLayout()
        other_form_layout.addWidget(QLabel("文件扩展名类型："), 0, 0)
        self.other_ext_combo = QComboBox()
        self.other_ext_combo.addItems([".txt (文本文件)", ".md (Markdown文档)", ".log (日志文件)", ".csv (表格数据)"])
        other_form_layout.addWidget(self.other_ext_combo, 0, 1)
        
        other_form_layout.addWidget(QLabel("自定义文件名："), 1, 0)
        self.other_name_edit = QLineEdit()
        self.other_name_edit.setPlaceholderText("留空则自动按时间戳命名")
        other_form_layout.addWidget(self.other_name_edit, 1, 1)
        tab_other_layout.addLayout(other_form_layout)
        
        tab_other_layout.addWidget(QLabel("文件初始详细描述："))
        self.other_content_edit = QTextEdit()
        self.other_content_edit.setPlaceholderText("在此输入文件初始化内容...")
        tab_other_layout.addWidget(self.other_content_edit)
        
        btn_gen_other = QPushButton("📝 生成其他文稿")
        btn_gen_other.clicked.connect(self.generate_other)
        tab_other_layout.addWidget(btn_gen_other)
        
        self.tab_widget.addTab(tab_other, "其他文件生成")
        
        # --- Tab 3: 档案封面生成 ---
        tab_cover = QWidget()
        tab_cover_layout = QVBoxLayout(tab_cover)
        
        cover_form_layout = QGridLayout()
        cover_form_layout.addWidget(QLabel("档案封面名称："), 0, 0)
        self.cover_name_edit = QLineEdit()
        self.cover_name_edit.setPlaceholderText("请输入档案封面主题名称")
        cover_form_layout.addWidget(self.cover_name_edit, 0, 1)
        
        cover_form_layout.addWidget(QLabel("全宗号 / 目录号："), 1, 0)
        self.archive_no_edit = QLineEdit()
        self.archive_no_edit.setPlaceholderText("例如: A-001")
        cover_form_layout.addWidget(self.archive_no_edit, 1, 1)
        
        cover_form_layout.addWidget(QLabel("保存期限："), 2, 0)
        self.archive_period_combo = QComboBox()
        self.archive_period_combo.addItems(["永久", "定期 30年", "定期 10年"])
        cover_form_layout.addWidget(self.archive_period_combo, 2, 1)
        
        tab_cover_layout.addLayout(cover_form_layout)
        
        tab_cover_layout.addWidget(QLabel("卷内备注与说明："))
        self.cover_desc_edit = QTextEdit()
        self.cover_desc_edit.setPlaceholderText("请输入档案卷盒封面的详细说明信息...")
        tab_cover_layout.addWidget(self.cover_desc_edit)
        
        btn_gen_cover = QPushButton("📑 生成档案封面说明")
        btn_gen_cover.clicked.connect(self.generate_cover)
        tab_cover_layout.addWidget(btn_gen_cover)
        
        self.tab_widget.addTab(tab_cover, "档案封面生成")
        
        # 底部实用按钮
        bottom_action_layout = QHBoxLayout()
        btn_open_current = QPushButton("📂 打开当前工作目录")
        btn_open_current.clicked.connect(self.open_current_dir)
        bottom_action_layout.addWidget(btn_open_current)
        right_layout.addLayout(bottom_action_layout)
        
        # 设置左右分栏等宽比例 (1:1)
        half_width = app_width // 2 - 20
        splitter.setSizes([half_width, half_width])
        
        # 加载历史保存路径或默认主目录
        saved_path = self.settings.value("last_path", QDir.homePath())
        if not os.path.exists(saved_path):
            saved_path = QDir.homePath()
        self.path_edit.setText(saved_path)
        self.update_tree_view(saved_path)
        
        # 状态栏
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        if DOCX_AVAILABLE:
            self.status_bar.showMessage("系统就绪 - 国标公文格式引擎已就绪")
        else:
            self.status_bar.showMessage("⚠️ 提示: 未检测到 python-docx 库，请使用 pip install python-docx 安装")
        
        # 初始化预览
        self.update_preview()
        
    def on_path_changed(self, text):
        if os.path.exists(text) and os.path.isdir(text):
            self.update_tree_view(text)
            
    def update_tree_view(self, path):
        if os.path.exists(path) and os.path.isdir(path):
            index = self.file_model.index(path)
            self.tree_view.setRootIndex(index)
            self.file_path_edit.setText(path)
            
    def browse_path(self):
        dir_path = QFileDialog.getExistingDirectory(self, "选择目标文件夹", self.path_edit.text())
        if dir_path:
            self.path_edit.setText(dir_path)
            self.set_current_path()
            
    def go_parent_path(self):
        current_path = self.path_edit.text().strip()
        if os.path.exists(current_path) and os.path.isdir(current_path):
            parent_path = os.path.dirname(current_path)
            if parent_path and os.path.exists(parent_path):
                self.path_edit.setText(parent_path)
                self.set_current_path()
            else:
                QMessageBox.information(self, "提示", "已经是盘符根目录或最上层目录！")
        else:
            QMessageBox.warning(self, "警告", "当前路径无效！")
            
    def set_current_path(self):
        current_path = self.path_edit.text().strip()
        if os.path.exists(current_path) and os.path.isdir(current_path):
            self.update_tree_view(current_path)
            self.settings.setValue("last_path", current_path)
            self.status_bar.showMessage(f"已锁定工作路径: {current_path}", 4000)
        else:
            QMessageBox.warning(self, "警告", "输入的路径无效或不存在！")
            
    def on_tree_double_click(self, index):
        file_path = self.file_model.filePath(index)
        if os.path.isdir(file_path):
            self.path_edit.setText(file_path)
            self.update_tree_view(file_path)
            
    def open_current_dir(self):
        current_path = self.get_current_path()
        if os.path.exists(current_path):
            try:
                if sys.platform == 'linux':
                    subprocess.Popen(['xdg-open', current_path])
                elif sys.platform == 'darwin':
                    subprocess.Popen(['open', current_path])
                else:
                    subprocess.Popen(['explorer', current_path])
            except Exception as e:
                QMessageBox.warning(self, "错误", f"无法唤起文件管理器：{str(e)}")
        else:
            QMessageBox.warning(self, "提示", "当前路径不存在！")
            
    def get_current_path(self):
        path = self.path_edit.text().strip()
        if not os.path.exists(path) or not os.path.isdir(path):
            path = QDir.homePath()
        return path
        
    def get_time_string(self):
        now = datetime.now()
        if self.time_format_combo.currentIndex() == 0:
            return now.strftime("%Y.%m.%d")
        else:
            return now.strftime("%Y.%m.%d %H%M%S")
            
    def update_preview(self):
        folder_name = self.folder_name_edit.text().strip()
        time_str = self.get_time_string()
        if folder_name:
            preview = f"{folder_name} {time_str}"
        else:
            preview = f"新文件夹 {time_str}"
        self.preview_label.setText(preview)
        
    def create_folder(self):
        folder_name = self.folder_name_edit.text().strip() or "新文件夹"
        full_name = f"{folder_name} {self.get_time_string()}"
        current_path = self.get_current_path()
        full_path = os.path.join(current_path, full_name)
        
        try:
            os.makedirs(full_path, exist_ok=True)
            self.status_bar.showMessage(f"✅ 成功创建文件夹: {full_name}", 4000)
            QMessageBox.information(self, "成功", f"主文件夹已成功创建：\n{full_path}")
            self.update_tree_view(current_path)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"创建文件夹失败：\n{str(e)}")
            
    def create_sub_folders(self):
        folder_name = self.folder_name_edit.text().strip() or "新文件夹"
        full_name = f"{folder_name} {self.get_time_string()}"
        current_path = self.get_current_path()
        full_path = os.path.join(current_path, full_name)
        
        sub_names = self.sub_folder_edit.text().strip()
        sub_list = [s.strip() for s in sub_names.split(',') if s.strip()]
        
        if not sub_list:
            QMessageBox.warning(self, "提示", "请输入至少一个子文件夹名称（用逗号分隔）！")
            return
            
        try:
            os.makedirs(full_path, exist_ok=True)
            for sub_name in sub_list:
                os.makedirs(os.path.join(full_path, sub_name), exist_ok=True)
                
            self.status_bar.showMessage(f"✅ 成功创建多级目录", 4000)
            QMessageBox.information(self, "成功", f"文件夹结构已成功创建：\n{full_path}")
            self.update_tree_view(current_path)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"创建失败：\n{str(e)}")
            
    def generate_document(self):
        """按照 GB/T 9704-2012 国家标准生成符合公文格式的 Word 文档"""
        if not DOCX_AVAILABLE:
            QMessageBox.critical(self, "错误", "未检测到 python-docx 库！\n请先在终端运行：pip install python-docx")
            return
            
        current_path = self.get_current_path()
        title = self.doc_title_edit.text().strip() or "公文标题"
        level = self.doc_level_combo.currentText()
        content = self.doc_content_edit.toPlainText()
        
        time_str = self.get_time_string()
        file_name = f"{title}_{time_str}.docx"
        file_path = os.path.join(current_path, file_name)
        
        try:
            doc = Document()
            
            # 1. 设置公文标题 (2号 方正小标宋简体，居中)
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_p.paragraph_format.space_before = Pt(12)
            title_p.paragraph_format.space_after = Pt(12)
            
            title_run = title_p.add_run(title)
            title_run.font.name = '方正小标宋简体'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
            title_run.font.size = Pt(22)  # 2号字对应 22pt
            title_run.font.bold = False
            
            # 2. 密级与元数据段落 (3号 黑体)
            p_meta = doc.add_paragraph()
            p_meta.paragraph_format.space_after = Pt(6)
            run_level = p_meta.add_run(f"密级与紧急程度：{level}\n")
            run_level.font.name = '黑体'
            run_level._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run_level.font.size = Pt(16)
            run_level.font.bold = True
            
            # 3. 正文内容设置 (3号 仿宋_GB2312，首行缩进2字符，行距30磅)
            lines = content.split('\n')
            for line in lines:
                if not line.strip():
                    continue
                
                p = doc.add_paragraph()
                if line.startswith("一、") or line.startswith("二、") or line.startswith("三、") or line.startswith("四、"):
                    # 一级标题：3号黑体
                    run = p.add_run(line)
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(16)
                    p.paragraph_format.left_indent = Pt(0)
                elif line.startswith("（一）") or line.startswith("（二）") or line.startswith("（三）"):
                    # 二级标题：3号楷体加粗
                    run = p.add_run(line)
                    run.font.name = '楷体_GB2312'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
                    run.font.size = Pt(16)
                    run.font.bold = True
                    p.paragraph_format.left_indent = Pt(0)
                else:
                    # 普通正文：3号仿宋，首行缩进2字符，30磅行距
                    run = p.add_run(line)
                    run.font.name = '仿宋_GB2312'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(16)
                    p.paragraph_format.first_line_indent = Pt(32)  # 仿宋3号字2个字符约等于32pt
                    p.paragraph_format.line_spacing = 30           # 固定行距 30 磅
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
            
            doc.save(file_path)
                
            self.status_bar.showMessage(f"✅ 已生成国标公文 Word 文档: {file_name}", 4000)
            QMessageBox.information(self, "成功", f"符合国标格式的公文文件已生成：\n{file_path}")
            self.update_tree_view(current_path)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成国标公文 Word 文档失败：\n{str(e)}")
            
    def generate_other(self):
        current_path = self.get_current_path()
        ext_idx = self.other_ext_combo.currentIndex()
        exts = [".txt", ".md", ".log", ".csv"]
        ext = exts[ext_idx]
        
        custom_name = self.other_name_edit.text().strip()
        time_str = self.get_time_string()
        
        if custom_name:
            file_name = f"{custom_name}_{time_str}{ext}"
        else:
            file_name = f"文稿_{time_str}{ext}"
            
        file_path = os.path.join(current_path, file_name)
        content = self.other_content_edit.toPlainText()
        
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
                
            self.status_bar.showMessage(f"✅ 已生成文件: {file_name}", 4000)
            QMessageBox.information(self, "成功", f"文件已生成：\n{file_path}")
            self.update_tree_view(current_path)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成失败：\n{str(e)}")
            
    def generate_cover(self):
        current_path = self.get_current_path()
        cover_name = self.cover_name_edit.text().strip() or "未命名档案"
        archive_no = self.archive_no_edit.text().strip() or "无"
        period = self.archive_period_combo.currentText()
        desc = self.cover_desc_edit.toPlainText()
        
        time_str = self.get_time_string()
        file_name = f"封面_{cover_name}_{time_str}.txt"
        file_path = os.path.join(current_path, file_name)
        
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write("【档案卷盒封面说明】\n")
                f.write(f"档案名称：{cover_name}\n")
                f.write(f"全宗/目录号：{archive_no}\n")
                f.write(f"保存期限：{period}\n")
                f.write(f"创建时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("=" * 50 + "\n\n")
                f.write(f"详细说明：\n{desc}")
                
            self.status_bar.showMessage(f"✅ 已生成档案封面说明: {file_name}", 4000)
            QMessageBox.information(self, "成功", f"档案封面说明已生成：\n{file_path}")
            self.update_tree_view(current_path)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成失败：\n{str(e)}")


def main():
    app = QApplication(sys.argv)
    window = GenDirApp()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
